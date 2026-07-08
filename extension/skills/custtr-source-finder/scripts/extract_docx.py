"""
Extract content from a .docx file for source-finding.
Outputs JSON with per-section text, organized by heading.

Usage:
    python extract_docx.py <path_to_docx>
"""

import sys
import re
import json
from docx import Document
from docx.oxml.ns import qn

_NOISE_PATTERNS = [
    re.compile(r"OST and Script", re.IGNORECASE),
    re.compile(r"Fully Shared", re.IGNORECASE),
    re.compile(r"^Course Name\s*:", re.IGNORECASE),
    re.compile(r"^Module Name\s*:", re.IGNORECASE),
    re.compile(r"The diagram is already created", re.IGNORECASE),
    re.compile(r"Please decide a presentation template", re.IGNORECASE),
    re.compile(r"^VO[\s«]", re.IGNORECASE),
    re.compile(r"^local instruction & data memory", re.IGNORECASE),
    re.compile(r"^RPU$", re.IGNORECASE),
    # Headers / footers / boilerplate
    re.compile(r"^(©|copyright|\(c\))\s", re.IGNORECASE),
    re.compile(r"all rights reserved", re.IGNORECASE),
    re.compile(r"^confidential\b", re.IGNORECASE),
    re.compile(r"^internal use only\b", re.IGNORECASE),
    re.compile(r"^(duration|learning path|level)\s*:", re.IGNORECASE),
    re.compile(r"page\s+\d+\s+of\s+\d+", re.IGNORECASE),
]


def is_noise(line):
    return any(p.search(line) for p in _NOISE_PATTERNS)


def is_heading(para):
    return para.style.name.startswith("Heading")


def heading_level(para):
    parts = para.style.name.split()
    try:
        return int(parts[-1])
    except (ValueError, IndexError):
        return 1


def has_image(para):
    return bool(para._element.findall(".//" + qn("a:blip")))


def para_text_no_strikethrough(para):
    """Return paragraph text with strikethrough runs excluded. Cached per call."""
    parts = []
    for run in para.runs:
        if run.font.strike:
            continue
        if run.text.strip():
            parts.append(run.text)
    return "".join(parts).strip()


def extract(docx_path):
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"Error reading DOCX: {e}. Ensure the file is a valid, unencrypted Word document.", file=sys.stderr)
        sys.exit(1)

    sections = []
    section_num = 0

    def new_section(heading, level):
        nonlocal section_num
        section_num += 1
        return {"section": section_num, "heading": heading, "level": level, "text": [], "has_images": False}

    current = new_section("(document start)", 0)

    for para in doc.paragraphs:
        raw_text = para.text.strip()
        clean_text = para_text_no_strikethrough(para)

        if is_heading(para) and raw_text:
            if current["text"] or current["has_images"]:
                sections.append(current)
            current = new_section(raw_text, heading_level(para))
        else:
            if has_image(para):
                current["has_images"] = True
            if clean_text and not is_noise(clean_text):
                current["text"].append(clean_text)

    if current["text"] or current["has_images"]:
        sections.append(current)

    # Tables — fix: cache para_text_no_strikethrough result to avoid double-call
    for table in doc.tables:
        table_texts = []
        for row in table.rows:
            for cell in row.cells:
                cleaned_paras = [para_text_no_strikethrough(p) for p in cell.paragraphs]
                cell_text = " ".join(p for p in cleaned_paras if p).strip()
                if cell_text and not is_noise(cell_text):
                    table_texts.append(cell_text)
        if table_texts:
            section_num += 1
            sections.append({
                "section": section_num,
                "heading": "(table)",
                "level": 0,
                "text": table_texts,
                "has_images": False,
            })

    return sections


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python extract_docx.py <path_to_docx>", file=sys.stderr)
        sys.exit(1)

    data = extract(sys.argv[1])
    sys.stdout.buffer.write(json.dumps(data, indent=2, ensure_ascii=True).encode("utf-8"))
    sys.stdout.buffer.write(b"\n")
